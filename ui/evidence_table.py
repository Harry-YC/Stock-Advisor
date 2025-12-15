"""
Evidence Table UI Component

Displays the extraction table with structured evidence data:
- PMID, Title, Design, N, Population, Intervention, Comparator, Outcomes, Effect, Quality
- Extract button per paper (calls EvidenceExtractor)
- Editable cells for human correction
- Export to CSV/Word
- Bulk extract for selected papers
"""

import streamlit as st
import pandas as pd
from datetime import datetime
from typing import List, Dict, Optional, Any
import io
import json

from core.evidence_corpus import (
    EvidenceCorpus,
    ExtractedEvidence,
    get_corpus_from_session,
    init_corpus_in_session
)
from config import settings


def render_evidence_table():
    """
    Render the main evidence extraction table.

    Shows all included papers with their extracted data and allows
    users to extract, edit, and verify evidence.
    """
    st.header("Evidence Extraction Table")

    # Get corpus
    corpus = init_corpus_in_session()

    # Stats bar
    stats = corpus.get_stats()
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Included Papers", stats["included"])
    with col2:
        st.metric("With Extractions", stats["with_extraction"])
    with col3:
        coverage_pct = f"{stats['extraction_coverage']:.0%}"
        st.metric("Extraction Coverage", coverage_pct)
    with col4:
        st.metric("Pending Review", stats["pending"])

    st.markdown("---")

    # Check if we have included papers
    if not corpus.included_pmids:
        st.info("""
        **No papers included yet.**

        Go to the Evidence Library tab to search for papers and screen them for inclusion.
        Once papers are included, they will appear here for evidence extraction.
        """)
        return

    # Action bar
    col_actions = st.columns([1, 1, 1, 2])
    with col_actions[0]:
        if st.button("Extract All Missing", type="primary"):
            _bulk_extract_missing(corpus)
    with col_actions[1]:
        if st.button("Export CSV"):
            _export_to_csv(corpus)
    with col_actions[2]:
        if st.button("Export Word"):
            _export_to_word(corpus)

    st.markdown("---")

    # Filter options
    with st.expander("Filters", expanded=False):
        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            filter_extracted = st.selectbox(
                "Extraction Status",
                ["All", "Extracted", "Not Extracted"],
                key="evidence_table_filter_extracted"
            )
        with col_f2:
            filter_verified = st.selectbox(
                "Verification Status",
                ["All", "Human Verified", "AI Only"],
                key="evidence_table_filter_verified"
            )
        with col_f3:
            filter_design = st.multiselect(
                "Study Design",
                ["RCT", "Cohort", "Case Series", "Case Report", "Systematic Review", "Other"],
                key="evidence_table_filter_design"
            )

    # Build table data
    table_data = _build_table_data(corpus, filter_extracted, filter_verified, filter_design)

    if not table_data:
        st.warning("No papers match the current filters.")
        return

    # Render table
    st.markdown(f"### Showing {len(table_data)} papers")

    # Use tabs for table vs card view
    tab_table, tab_cards = st.tabs(["Table View", "Card View"])

    with tab_table:
        _render_table_view(corpus, table_data)

    with tab_cards:
        _render_card_view(corpus, table_data)


def _build_table_data(
    corpus: EvidenceCorpus,
    filter_extracted: str,
    filter_verified: str,
    filter_design: List[str]
) -> List[Dict]:
    """Build table data from corpus with filters applied."""
    data = []

    for pmid in corpus.included_pmids:
        extraction = corpus.get_extraction(pmid)

        # Apply filters
        if filter_extracted == "Extracted" and not extraction:
            continue
        if filter_extracted == "Not Extracted" and extraction:
            continue
        if filter_verified == "Human Verified" and (not extraction or not extraction.human_verified):
            continue
        if filter_verified == "AI Only" and extraction and extraction.human_verified:
            continue
        if filter_design and extraction and extraction.study_design not in filter_design:
            continue

        # Build row
        if extraction:
            row = extraction.to_table_row()
            row["has_extraction"] = True
            row["human_verified"] = extraction.human_verified
        else:
            row = {
                "PMID": pmid,
                "Title": _get_title_from_session(pmid),
                "Design": "",
                "N": "",
                "Population": "",
                "Intervention": "",
                "Comparator": "",
                "Outcomes": "",
                "Effect": "",
                "Verified": "No",
                "has_extraction": False,
                "human_verified": False
            }

        data.append(row)

    return data


def _get_title_from_session(pmid: str) -> str:
    """Get paper title from session state search results."""
    search_results = st.session_state.get('search_results', [])
    if isinstance(search_results, dict):
        search_results = search_results.get('citations', [])

    for cit in search_results:
        cit_pmid = cit.pmid if hasattr(cit, 'pmid') else cit.get('pmid', '')
        if str(cit_pmid) == str(pmid):
            return cit.title if hasattr(cit, 'title') else cit.get('title', f'PMID {pmid}')

    return f"PMID {pmid}"


def _render_table_view(corpus: EvidenceCorpus, table_data: List[Dict]):
    """Render the data as a table."""
    # Create DataFrame for display
    df = pd.DataFrame(table_data)

    # Select display columns
    display_cols = ["PMID", "Title", "Design", "N", "Population", "Intervention", "Comparator", "Outcomes", "Effect", "Verified"]
    df_display = df[display_cols]

    # Style the dataframe
    st.dataframe(
        df_display,
        use_container_width=True,
        hide_index=True,
        column_config={
            "PMID": st.column_config.TextColumn("PMID", width="small"),
            "Title": st.column_config.TextColumn("Title", width="medium"),
            "Design": st.column_config.TextColumn("Design", width="small"),
            "N": st.column_config.TextColumn("N", width="small"),
            "Population": st.column_config.TextColumn("Population", width="medium"),
            "Intervention": st.column_config.TextColumn("Intervention", width="medium"),
            "Comparator": st.column_config.TextColumn("Comparator", width="medium"),
            "Outcomes": st.column_config.TextColumn("Outcomes", width="medium"),
            "Effect": st.column_config.TextColumn("Effect", width="small"),
            "Verified": st.column_config.TextColumn("Verified", width="small"),
        }
    )

    # Per-row actions below the table
    st.markdown("### Paper Actions")
    selected_pmid = st.selectbox(
        "Select paper to edit/extract",
        [row["PMID"] for row in table_data],
        key="evidence_table_select_pmid"
    )

    if selected_pmid:
        _render_paper_actions(corpus, selected_pmid)


def _render_card_view(corpus: EvidenceCorpus, table_data: List[Dict]):
    """Render the data as expandable cards."""
    for row in table_data:
        pmid = row["PMID"]
        has_extraction = row.get("has_extraction", False)
        human_verified = row.get("human_verified", False)

        # Status indicator
        if has_extraction and human_verified:
            status_icon = "🟢"
            status_text = "Verified"
        elif has_extraction:
            status_icon = "🟡"
            status_text = "AI Extracted"
        else:
            status_icon = "🔴"
            status_text = "Not Extracted"

        with st.expander(f"{status_icon} PMID {pmid}: {row.get('Title', '')[:50]}...", expanded=False):
            if has_extraction:
                # Show extraction data
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**Study Design:** {row.get('Design', 'NR')}")
                    st.markdown(f"**Sample Size:** {row.get('N', 'NR')}")
                    st.markdown(f"**Population:** {row.get('Population', 'NR')}")
                with col2:
                    st.markdown(f"**Intervention:** {row.get('Intervention', 'NR')}")
                    st.markdown(f"**Comparator:** {row.get('Comparator', 'NR')}")
                    st.markdown(f"**Effect:** {row.get('Effect', 'NR')}")

                st.markdown(f"**Outcomes:** {row.get('Outcomes', 'NR')}")

                # Edit button
                if st.button(f"Edit Extraction", key=f"edit_{pmid}"):
                    st.session_state[f"editing_extraction_{pmid}"] = True
                    st.rerun()

                # Verify button
                if not human_verified:
                    if st.button(f"Mark as Verified", key=f"verify_{pmid}"):
                        extraction = corpus.get_extraction(pmid)
                        if extraction:
                            extraction.human_verified = True
                            extraction.verification_notes = f"Verified by user on {datetime.now().strftime('%Y-%m-%d')}"
                            corpus.add_extraction(extraction)
                            st.success("Marked as verified!")
                            st.rerun()
            else:
                st.info("No extraction data yet.")
                if st.button(f"Extract Now", key=f"extract_{pmid}"):
                    _extract_single_paper(corpus, pmid)


def _render_paper_actions(corpus: EvidenceCorpus, pmid: str):
    """Render action buttons and forms for a single paper."""
    extraction = corpus.get_extraction(pmid)

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("Extract / Re-extract", key=f"action_extract_{pmid}"):
            _extract_single_paper(corpus, pmid)

    with col2:
        if extraction:
            if st.button("Edit Extraction", key=f"action_edit_{pmid}"):
                st.session_state[f"editing_extraction_{pmid}"] = True
                st.rerun()

    with col3:
        if extraction and not extraction.human_verified:
            if st.button("Verify", key=f"action_verify_{pmid}"):
                extraction.human_verified = True
                corpus.add_extraction(extraction)
                st.success("Verified!")
                st.rerun()

    # Edit form
    if st.session_state.get(f"editing_extraction_{pmid}", False):
        st.markdown("---")
        st.markdown("### Edit Extraction")
        _render_extraction_edit_form(corpus, pmid, extraction)


def _render_extraction_edit_form(corpus: EvidenceCorpus, pmid: str, extraction: Optional[ExtractedEvidence]):
    """Render form for editing extraction data."""
    with st.form(f"edit_form_{pmid}"):
        col1, col2 = st.columns(2)

        with col1:
            study_design = st.selectbox(
                "Study Design",
                ["RCT", "Cohort", "Case-Control", "Case Series", "Case Report", "Systematic Review", "Meta-Analysis", "Other"],
                index=0 if not extraction else ["RCT", "Cohort", "Case-Control", "Case Series", "Case Report", "Systematic Review", "Meta-Analysis", "Other"].index(extraction.study_design) if extraction.study_design in ["RCT", "Cohort", "Case-Control", "Case Series", "Case Report", "Systematic Review", "Meta-Analysis", "Other"] else 7,
                key=f"edit_design_{pmid}"
            )
            sample_size = st.number_input(
                "Sample Size",
                min_value=0,
                value=extraction.sample_size if extraction and extraction.sample_size else 0,
                key=f"edit_n_{pmid}"
            )
            population = st.text_area(
                "Population",
                value=extraction.population if extraction else "",
                key=f"edit_pop_{pmid}"
            )

        with col2:
            intervention = st.text_area(
                "Intervention",
                value=extraction.intervention if extraction else "",
                key=f"edit_int_{pmid}"
            )
            comparator = st.text_area(
                "Comparator",
                value=extraction.comparator if extraction else "",
                key=f"edit_comp_{pmid}"
            )
            effect_size = st.text_input(
                "Effect Size",
                value=extraction.effect_size if extraction and extraction.effect_size else "",
                key=f"edit_effect_{pmid}"
            )

        outcomes = st.text_area(
            "Outcomes (comma-separated)",
            value=", ".join(extraction.outcomes) if extraction else "",
            key=f"edit_outcomes_{pmid}"
        )

        key_findings = st.text_area(
            "Key Findings",
            value=extraction.key_findings if extraction else "",
            key=f"edit_findings_{pmid}"
        )

        col_submit = st.columns([1, 1, 2])
        with col_submit[0]:
            submitted = st.form_submit_button("Save", type="primary")
        with col_submit[1]:
            cancelled = st.form_submit_button("Cancel")

        if submitted:
            new_extraction = ExtractedEvidence(
                pmid=pmid,
                title=extraction.title if extraction else _get_title_from_session(pmid),
                study_design=study_design,
                sample_size=sample_size if sample_size > 0 else None,
                population=population,
                intervention=intervention,
                comparator=comparator,
                outcomes=[o.strip() for o in outcomes.split(",") if o.strip()],
                key_findings=key_findings,
                effect_size=effect_size if effect_size else None,
                extracted_by="human",
                human_verified=True
            )
            corpus.add_extraction(new_extraction)
            st.session_state[f"editing_extraction_{pmid}"] = False
            st.success("Extraction saved!")
            st.rerun()

        if cancelled:
            st.session_state[f"editing_extraction_{pmid}"] = False
            st.rerun()


def _extract_single_paper(corpus: EvidenceCorpus, pmid: str):
    """Extract evidence from a single paper using AI."""
    with st.spinner(f"Extracting evidence from PMID {pmid}..."):
        try:
            # Get abstract from session state
            abstract = _get_abstract_from_session(pmid)
            title = _get_title_from_session(pmid)

            if not abstract:
                st.warning(f"No abstract available for PMID {pmid}. Manual extraction required.")
                return

            # Call extraction service
            extraction = _call_extraction_service(pmid, title, abstract)

            if extraction:
                corpus.add_extraction(extraction)
                st.success(f"Extraction complete for PMID {pmid}")
                st.rerun()
            else:
                st.error("Extraction failed. Please try manual extraction.")

        except Exception as e:
            st.error(f"Extraction error: {e}")


def _get_abstract_from_session(pmid: str) -> str:
    """Get paper abstract from session state."""
    search_results = st.session_state.get('search_results', [])
    if isinstance(search_results, dict):
        search_results = search_results.get('citations', [])

    for cit in search_results:
        cit_pmid = cit.pmid if hasattr(cit, 'pmid') else cit.get('pmid', '')
        if str(cit_pmid) == str(pmid):
            return cit.abstract if hasattr(cit, 'abstract') else cit.get('abstract', '')

    return ""


def _call_extraction_service(pmid: str, title: str, abstract: str) -> Optional[ExtractedEvidence]:
    """Call the AI extraction service."""
    try:
        from core.data_extractor import EvidenceExtractor

        extractor = EvidenceExtractor(api_key=settings.OPENAI_API_KEY)
        result = extractor.extract_from_abstract(pmid, title, abstract)

        if result:
            return ExtractedEvidence(
                pmid=pmid,
                title=title,
                study_design=result.get('study_design', ''),
                population=result.get('population', ''),
                sample_size=result.get('sample_size'),
                intervention=result.get('intervention', ''),
                comparator=result.get('comparator', ''),
                outcomes=result.get('outcomes', []),
                key_findings=result.get('key_findings', ''),
                effect_size=result.get('effect_size'),
                extracted_by="ai",
                human_verified=False
            )
    except Exception as e:
        st.error(f"Extraction service error: {e}")
    return None


def _bulk_extract_missing(corpus: EvidenceCorpus):
    """Extract evidence for all papers that don't have extractions."""
    missing = corpus.get_unextracted_pmids()

    if not missing:
        st.info("All papers already have extractions.")
        return

    progress = st.progress(0)
    status = st.empty()

    extracted = 0
    failed = 0

    for i, pmid in enumerate(missing):
        status.text(f"Extracting {i+1}/{len(missing)}: PMID {pmid}")
        progress.progress((i + 1) / len(missing))

        try:
            abstract = _get_abstract_from_session(pmid)
            title = _get_title_from_session(pmid)

            if abstract:
                extraction = _call_extraction_service(pmid, title, abstract)
                if extraction:
                    corpus.add_extraction(extraction)
                    extracted += 1
                else:
                    failed += 1
            else:
                failed += 1
        except Exception:
            failed += 1

    progress.empty()
    status.empty()
    st.success(f"Bulk extraction complete: {extracted} extracted, {failed} failed")
    st.rerun()


def _export_to_csv(corpus: EvidenceCorpus):
    """Export evidence table to CSV."""
    extractions = corpus.get_citable_evidence()

    if not extractions:
        st.warning("No extractions to export.")
        return

    rows = [e.to_table_row() for e in extractions]
    df = pd.DataFrame(rows)

    csv = df.to_csv(index=False)
    st.download_button(
        label="Download CSV",
        data=csv,
        file_name=f"evidence_table_{datetime.now().strftime('%Y%m%d')}.csv",
        mime="text/csv"
    )


def _export_to_word(corpus: EvidenceCorpus):
    """Export evidence table to Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        st.error("python-docx not installed. Run: pip install python-docx")
        return

    extractions = corpus.get_citable_evidence()

    if not extractions:
        st.warning("No extractions to export.")
        return

    doc = Document()
    doc.add_heading('Evidence Extraction Table', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")}')
    doc.add_paragraph(f'Papers: {len(extractions)}')

    # Create table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    # Header row
    headers = ['PMID', 'Design', 'N', 'Population', 'Intervention', 'Comparator', 'Outcomes', 'Effect']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header

    # Data rows
    for ext in extractions:
        row = table.add_row().cells
        row[0].text = ext.pmid
        row[1].text = ext.study_design or ''
        row[2].text = str(ext.sample_size) if ext.sample_size else 'NR'
        row[3].text = ext.population[:50] if ext.population else ''
        row[4].text = ext.intervention[:50] if ext.intervention else ''
        row[5].text = ext.comparator[:50] if ext.comparator else ''
        row[6].text = ', '.join(ext.outcomes[:3]) if ext.outcomes else ''
        row[7].text = ext.effect_size or 'NR'

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Word",
        data=buffer.getvalue(),
        file_name=f"evidence_table_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def render_evidence_mini_table(max_rows: int = 5):
    """
    Render a mini version of the evidence table for display in other views.

    Args:
        max_rows: Maximum number of rows to show
    """
    corpus = get_corpus_from_session()
    if not corpus:
        return

    extractions = corpus.get_citable_evidence()[:max_rows]

    if not extractions:
        st.caption("No evidence extracted yet.")
        return

    st.markdown("**Included Evidence:**")
    for ext in extractions:
        st.caption(f"• PMID {ext.pmid}: {ext.study_design or 'Unknown'}, n={ext.sample_size or 'NR'}")
