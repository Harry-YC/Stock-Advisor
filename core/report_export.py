"""
Report Export Module

Generates full research reports in Word (.docx) format including:
- Research question and type
- Recommendation with confidence level
- Key findings (bullet points)
- Expert perspectives (expanded)
- Citations (formatted bibliography)

Compatible with: python-docx library
"""

from dataclasses import dataclass
from typing import List, Dict, Optional, Union
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.pubmed_client import Citation


@dataclass
class ReportConfig:
    """Configuration for report generation."""
    include_recommendation: bool = True
    include_key_findings: bool = True
    include_expert_perspectives: bool = True
    include_citations: bool = True
    include_validation_summary: bool = True
    include_dissenting_views: bool = True
    max_citations: int = 50
    citation_format: str = "vancouver"  # "vancouver" or "apa"


def generate_research_report(
    result: Union[Dict, 'ResearchResult'],
    config: Optional[ReportConfig] = None,
    project_name: Optional[str] = None
) -> BytesIO:
    """
    Generate a Word document (.docx) report from research results.

    Args:
        result: ResearchResult object or dict from st.session_state.research_result
        config: Optional ReportConfig for customization
        project_name: Optional project name for header

    Returns:
        BytesIO buffer containing the .docx file
    """
    config = config or ReportConfig()
    doc = Document()

    # Extract data (handle both dict and ResearchResult)
    if isinstance(result, dict):
        question = result.get('question', '')
        question_type = result.get('question_type', 'research')
        recommendation = result.get('recommendation', '')
        confidence = result.get('confidence', 'MEDIUM')
        key_findings = result.get('key_findings', [])
        expert_responses = result.get('expert_responses', {})
        evidence_summary = result.get('evidence_summary', {})
        validations = result.get('validations', {})
        dissenting_views = result.get('dissenting_views', [])
        metadata = result.get('metadata', {})
    else:
        question = getattr(result, 'question', '')
        question_type = getattr(result, 'question_type', 'research')
        recommendation = getattr(result, 'recommendation', '')
        confidence = getattr(result, 'confidence', 'MEDIUM')
        key_findings = getattr(result, 'key_findings', [])
        expert_responses = getattr(result, 'expert_responses', {})
        evidence_summary = getattr(result, 'evidence_summary', {})
        validations = getattr(result, 'validations', {})
        dissenting_views = getattr(result, 'dissenting_views', [])
        metadata = getattr(result, 'metadata', {})

    # Get citations from evidence_summary
    citations = []
    if evidence_summary:
        citations = evidence_summary.get('citations', []) if isinstance(evidence_summary, dict) else getattr(evidence_summary, 'citations', [])

    # Title page
    _add_title_page(doc, "Research Report", project_name or "Palliative Surgery GDG",
                    datetime.now().strftime("%B %d, %Y"))

    # Research Question section
    doc.add_page_break()
    _add_section_header(doc, "Research Question", level=1)

    # Question type badge
    type_name = _get_question_type_name(question_type)
    p = doc.add_paragraph()
    run = p.add_run(f"Question Type: {type_name}")
    run.bold = True
    run.font.size = Pt(11)

    # Question text
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(question).italic = True

    # Recommendation section
    if config.include_recommendation and recommendation:
        doc.add_paragraph()
        _add_section_header(doc, "Recommendation", level=1)

        # Confidence indicator
        p = doc.add_paragraph()
        conf_run = p.add_run(f"Confidence Level: {confidence}")
        conf_run.bold = True
        if confidence == "HIGH":
            conf_run.font.color.rgb = RGBColor(40, 167, 69)  # Green
        elif confidence == "LOW":
            conf_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
        else:
            conf_run.font.color.rgb = RGBColor(255, 193, 7)  # Yellow/Orange

        doc.add_paragraph()
        doc.add_paragraph(recommendation)

    # Key Findings section
    if config.include_key_findings and key_findings:
        doc.add_paragraph()
        _add_section_header(doc, "Key Findings", level=1)
        _add_bullet_list(doc, key_findings)

    # Validation Summary
    if config.include_validation_summary and validations:
        doc.add_paragraph()
        _add_section_header(doc, "Literature Validation Summary", level=1)

        total_supported = 0
        total_contradicted = 0
        total_no_evidence = 0

        for val in validations.values():
            if hasattr(val, 'claims_supported'):
                total_supported += val.claims_supported
                total_contradicted += val.claims_contradicted
                total_no_evidence += getattr(val, 'claims_no_evidence', 0)
            elif isinstance(val, dict):
                total_supported += val.get('claims_supported', 0)
                total_contradicted += val.get('claims_contradicted', 0)
                total_no_evidence += val.get('claims_no_evidence', 0)

        p = doc.add_paragraph()
        p.add_run(f"Claims supported by literature: {total_supported}").bold = True
        doc.add_paragraph(f"Claims contradicted by literature: {total_contradicted}")
        doc.add_paragraph(f"Claims with no available evidence: {total_no_evidence}")

    # Dissenting Views
    if config.include_dissenting_views and dissenting_views:
        doc.add_paragraph()
        _add_section_header(doc, "Dissenting Views", level=1)
        doc.add_paragraph("The following experts raised concerns or dissenting opinions:")
        _add_bullet_list(doc, dissenting_views)

    # Expert Perspectives section
    if config.include_expert_perspectives and expert_responses:
        doc.add_paragraph()
        _add_section_header(doc, "Expert Perspectives", level=1)

        for expert_name, response in expert_responses.items():
            _add_section_header(doc, expert_name, level=2)
            content = response.get('content', '') if isinstance(response, dict) else str(response)
            doc.add_paragraph(content)

    # Citations section
    if config.include_citations and citations:
        doc.add_paragraph()
        _add_section_header(doc, "References", level=1)
        _add_citations_section(doc, citations, config.citation_format, config.max_citations)

    # Metadata footer
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Generated by Palliative Surgery GDG on ").font.size = Pt(9)
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S")).font.size = Pt(9)
    elapsed = metadata.get('elapsed_seconds', 0) if isinstance(metadata, dict) else getattr(metadata, 'elapsed_seconds', 0)
    if elapsed:
        p.add_run(f" (Analysis time: {elapsed:.0f}s)").font.size = Pt(9)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_literature_report(
    citations: List[Union[Citation, Dict]],
    query: Optional[str] = None,
    project_name: Optional[str] = None,
    include_abstracts: bool = False
) -> BytesIO:
    """
    Generate a Word document with citation list from search results.

    Args:
        citations: List of Citation objects or dicts
        query: Optional search query for header
        project_name: Optional project name
        include_abstracts: Whether to include abstracts

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # Title
    _add_title_page(doc, "Literature Search Results", project_name or "Palliative Surgery GDG",
                    datetime.now().strftime("%B %d, %Y"))

    doc.add_page_break()

    # Query info
    if query:
        _add_section_header(doc, "Search Query", level=1)
        doc.add_paragraph(query)
        doc.add_paragraph()

    # Statistics
    _add_section_header(doc, "Summary", level=1)
    doc.add_paragraph(f"Total citations: {len(citations)}")
    doc.add_paragraph()

    # Citations
    _add_section_header(doc, "Citations", level=1)

    for i, cit in enumerate(citations, 1):
        # Format citation
        formatted = format_citation_vancouver(cit, i)
        p = doc.add_paragraph()
        p.add_run(formatted)

        if include_abstracts:
            abstract = cit.abstract if hasattr(cit, 'abstract') else cit.get('abstract', '')
            if abstract:
                p = doc.add_paragraph()
                p.add_run("Abstract: ").bold = True
                p.add_run(abstract[:1000])
                if len(abstract) > 1000:
                    p.add_run("...")
                p.paragraph_format.left_indent = Inches(0.5)

        doc.add_paragraph()  # Spacing

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Generated by Palliative Surgery GDG on ").font.size = Pt(9)
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S")).font.size = Pt(9)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def format_citation_vancouver(citation: Union[Citation, Dict], index: int) -> str:
    """Format citation in Vancouver style (numbered)."""
    if hasattr(citation, 'authors'):
        authors = citation.authors
        title = citation.title
        journal = citation.journal
        year = citation.year
        pmid = citation.pmid
        doi = citation.doi
    else:
        authors = citation.get('authors', [])
        title = citation.get('title', 'Untitled')
        journal = citation.get('journal', '')
        year = citation.get('year', '')
        pmid = citation.get('pmid', '')
        doi = citation.get('doi', '')

    # Format authors
    if authors:
        if isinstance(authors, list):
            if len(authors) > 6:
                author_str = ', '.join(authors[:6]) + ', et al.'
            else:
                author_str = ', '.join(authors)
        else:
            author_str = str(authors)
    else:
        author_str = 'Unknown'

    # Build citation string
    parts = [f"{index}. {author_str}. {title}."]
    if journal:
        parts.append(f" {journal}.")
    if year:
        parts.append(f" {year}.")
    if pmid:
        parts.append(f" PMID: {pmid}.")
    if doi:
        parts.append(f" DOI: {doi}.")

    return ''.join(parts)


def format_citation_apa(citation: Union[Citation, Dict]) -> str:
    """Format citation in APA style."""
    if hasattr(citation, 'authors'):
        authors = citation.authors
        title = citation.title
        journal = citation.journal
        year = citation.year
    else:
        authors = citation.get('authors', [])
        title = citation.get('title', 'Untitled')
        journal = citation.get('journal', '')
        year = citation.get('year', 'n.d.')

    # Format authors (APA style)
    if authors:
        if isinstance(authors, list):
            if len(authors) == 1:
                author_str = authors[0]
            elif len(authors) == 2:
                author_str = f"{authors[0]} & {authors[1]}"
            elif len(authors) <= 7:
                author_str = ', '.join(authors[:-1]) + f', & {authors[-1]}'
            else:
                author_str = ', '.join(authors[:6]) + ', ... ' + authors[-1]
        else:
            author_str = str(authors)
    else:
        author_str = 'Unknown'

    return f"{author_str} ({year}). {title}. {journal}."


def _get_question_type_name(question_type: str) -> str:
    """Get display name for question type."""
    # Import here to avoid circular imports
    try:
        from core.question_templates import QUESTION_TYPES
        type_info = QUESTION_TYPES.get(question_type, {})
        return type_info.get('name', question_type.replace('_', ' ').title())
    except ImportError:
        return question_type.replace('_', ' ').title()


def _add_title_page(doc: Document, title: str, project_name: str, date: str):
    """Add a title page to the document."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Project name
    p = doc.add_paragraph()
    run = p.add_run(project_name)
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Date
    p = doc.add_paragraph()
    p.add_run(date).font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_section_header(doc: Document, text: str, level: int = 1):
    """Add a section header with consistent styling."""
    if level == 1:
        doc.add_heading(text, level=1)
    elif level == 2:
        doc.add_heading(text, level=2)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)


def _add_bullet_list(doc: Document, items: List[str]):
    """Add a bullet list to the document."""
    for item in items:
        doc.add_paragraph(str(item), style='List Bullet')


def _add_citations_section(
    doc: Document,
    citations: List[Union[Citation, Dict]],
    format_style: str = "vancouver",
    max_citations: int = 50
):
    """Add formatted citations section."""
    for i, cit in enumerate(citations[:max_citations], 1):
        if format_style == "vancouver":
            formatted = format_citation_vancouver(cit, i)
        else:
            formatted = format_citation_apa(cit)

        p = doc.add_paragraph()
        p.add_run(formatted)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    if len(citations) > max_citations:
        p = doc.add_paragraph()
        p.add_run(f"... and {len(citations) - max_citations} more citations").italic = True
