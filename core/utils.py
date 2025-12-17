"""
Core Utilities Module

Contains shared helper functions used across the application.
"""

import re
import logging
from typing import Any, Optional

logger = logging.getLogger(__name__)

__all__ = [
    'extract_simple_query',
    'user_friendly_error',
    'extract_text_from_file',
]


# =============================================================================
# ERROR HANDLING
# =============================================================================

# Mapping of exception types to user-friendly messages
ERROR_MESSAGES = {
    # Network/API errors
    'ConnectionError': "Unable to connect to the service. Please check your internet connection.",
    'TimeoutError': "The request took too long. Please try again.",
    'HTTPError': "The service returned an error. Please try again later.",
    'RateLimitError': "Too many requests. Please wait a moment and try again.",
    'AuthenticationError': "Authentication failed. Please check your API key.",
    'APIError': "The AI service encountered an issue. Please try again.",

    # File processing errors
    'FileNotFoundError': "The file could not be found.",
    'PermissionError': "Permission denied. Cannot access the file.",
    'UnicodeDecodeError': "Unable to read the file. It may be in an unsupported format.",
    'pdf': "Unable to process PDF. The file may be corrupted or password-protected.",

    # Data processing errors
    'JSONDecodeError': "Unable to parse the response. Please try again.",
    'KeyError': "Missing expected data in response.",
    'ValueError': "Invalid data format received.",

    # Generic fallbacks
    'default': "An unexpected error occurred. Please try again.",
}


def user_friendly_error(e: Exception, context: str = "") -> str:
    """
    Convert an exception to a user-friendly error message.

    Logs the full exception for debugging but returns a clean message for users.

    Args:
        e: The exception that occurred
        context: Optional context about what operation failed (e.g., "processing document")

    Returns:
        A user-friendly error message string
    """
    # Log the full exception for debugging
    logger.error(f"Error in {context}: {type(e).__name__}: {str(e)}", exc_info=True)

    # Get the exception type name
    exc_type = type(e).__name__
    exc_str = str(e).lower()

    # Try to match specific error types
    if exc_type in ERROR_MESSAGES:
        message = ERROR_MESSAGES[exc_type]
    # Check for known error patterns in the message
    elif 'rate limit' in exc_str or 'quota' in exc_str:
        message = ERROR_MESSAGES['RateLimitError']
    elif 'timeout' in exc_str:
        message = ERROR_MESSAGES['TimeoutError']
    elif 'connection' in exc_str or 'network' in exc_str:
        message = ERROR_MESSAGES['ConnectionError']
    elif 'authentication' in exc_str or 'api key' in exc_str or 'unauthorized' in exc_str:
        message = ERROR_MESSAGES['AuthenticationError']
    elif 'pdf' in exc_str:
        message = ERROR_MESSAGES['pdf']
    else:
        message = ERROR_MESSAGES['default']

    # Add context if provided
    if context:
        return f"{message} (while {context})"
    return message

def extract_simple_query(original_query: str) -> str:
    """
    Extract core medical terms from natural language query for fallback search.

    Creates a simpler, less restrictive query when AI-optimized query returns 0 results.
    """
    # Common stop words to remove
    stop_words = {
        'in', 'with', 'how', 'do', 'does', 'what', 'when', 'where', 'for',
        'and', 'or', 'the', 'a', 'an', 'to', 'of', 'are', 'is', 'was', 'were',
        'have', 'has', 'had', 'compare', 'versus', 'vs', 'between', 'among'
    }

    # Remove question marks and normalize
    text = original_query.lower().replace('?', '').replace(',', ' ')

    # Extract quoted phrases first (these are important)
    quoted = re.findall(r'"([^"]+)"', text)

    # Split into words
    words = text.split()

    # Keep capitalized acronyms from original (GOO, SEMS, etc.)
    acronyms = re.findall(r'\b[A-Z]{2,}\b', original_query)

    # Filter words: keep medical terms, remove stop words
    key_terms = []
    current_phrase = []

    for word in words:
        word_clean = re.sub(r'[^\w\s-]', '', word)  # Remove punctuation except hyphens

        if word_clean in stop_words or len(word_clean) < 3:
            if current_phrase:
                key_terms.append(' '.join(current_phrase))
                current_phrase = []
            continue

        current_phrase.append(word_clean)

    if current_phrase:
        key_terms.append(' '.join(current_phrase))

    # Combine quoted phrases, key terms, and acronyms
    all_terms = quoted + key_terms + [a.lower() for a in acronyms]

    # Remove duplicates while preserving order
    seen = set()
    unique_terms = []
    for term in all_terms:
        if term not in seen and term:
            seen.add(term)
            unique_terms.append(term)

    # Build simple query: use first 3-4 core terms with AND
    core_terms = unique_terms[:4]  # Limit to most important terms

    if len(core_terms) >= 2:
        # Multiple terms - use AND to combine
        simple_query = ' AND '.join([f'"{term}"' for term in core_terms])
    elif core_terms:
        # Single term
        simple_query = f'"{core_terms[0]}"'
    else:
        # Fallback: use first few words
        simple_query = ' '.join(original_query.split()[:5])

    return simple_query


# =============================================================================
# DOCUMENT EXTRACTION
# =============================================================================

def extract_text_from_file(file_path: str, mime_type: str = None) -> str:
    """
    Extract text content from uploaded files.

    Supports: PDF, Word (.docx), Excel (.xlsx), Text (.txt)

    Args:
        file_path: Path to the uploaded file
        mime_type: Optional MIME type hint

    Returns:
        Extracted text content
    """
    import os

    # Determine file type from extension if mime not provided
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == '.pdf' or (mime_type and 'pdf' in mime_type):
            return _extract_pdf(file_path)
        elif ext == '.docx' or (mime_type and 'wordprocessingml' in mime_type):
            return _extract_docx(file_path)
        elif ext == '.xlsx' or (mime_type and 'spreadsheetml' in mime_type):
            return _extract_xlsx(file_path)
        elif ext == '.txt' or (mime_type and 'text/plain' in mime_type):
            return _extract_text(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return f"Unsupported file type: {ext}"
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return f"Error reading file: {str(e)}"


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    text_parts = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"--- Page {i+1} ---\n{page_text}")

    return "\n\n".join(text_parts) if text_parts else "No text found in PDF"


def _extract_docx(file_path: str) -> str:
    """Extract text from Word document."""
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs) if paragraphs else "No text found in document"


def _extract_xlsx(file_path: str) -> str:
    """Extract text from Excel file."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_content = [f"=== Sheet: {sheet_name} ==="]

        for row in sheet.iter_rows():
            row_values = []
            for cell in row:
                if cell.value is not None:
                    row_values.append(str(cell.value))
            if row_values:
                sheet_content.append(" | ".join(row_values))

        if len(sheet_content) > 1:  # More than just header
            text_parts.append("\n".join(sheet_content))

    return "\n\n".join(text_parts) if text_parts else "No data found in spreadsheet"


def _extract_text(file_path: str) -> str:
    """Extract text from plain text file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()
