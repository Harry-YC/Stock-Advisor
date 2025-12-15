"""
Export UI Module for Palliative Surgery GDG

Provides comprehensive export functionality for:
- GDG discussion summaries
- Evidence tables
- Structured recommendations
- Full guideline documents
"""

import streamlit as st
from datetime import datetime
from io import BytesIO
from typing import Dict, List, Any, Optional

from config import settings


def render_export():
    """Render the Export tab interface."""
    st.header("Export")
    st.caption("Export GDG discussions, evidence, and recommendations")

    # Check what data is available
    research_result = st.session_state.get('research_result')
    expert_discussion = st.session_state.get('expert_discussion', {})
    search_results = st.session_state.get('search_results', {})
    cdp_sections = st.session_state.get('cdp_sections', {})

    # Summary of available content
    _render_content_summary(research_result, expert_discussion, search_results, cdp_sections)

    if not research_result and not expert_discussion and not search_results:
        st.info("No content available to export. Ask a question or run a GDG discussion first.")
        return

    st.markdown("---")

    # Export format selection
    col1, col2 = st.columns([1, 1])
    with col1:
        export_format = st.selectbox(
            "Export Format",
            ["Markdown", "Word Document", "JSON"],
            key="export_format_select"
        )

    with col2:
        export_scope = st.selectbox(
            "Export Scope",
            _get_available_scopes(research_result, expert_discussion, search_results, cdp_sections),
            key="export_scope_select"
        )

    # Content selection
    st.markdown("### Content to Include")
    col1, col2 = st.columns(2)

    with col1:
        include_question = st.checkbox("Clinical Question", value=True, key="export_include_question")
        include_evidence = st.checkbox("Evidence Table", value=True, key="export_include_evidence")
        include_recommendation = st.checkbox("Recommendation", value=True, key="export_include_recommendation")

    with col2:
        include_discussion = st.checkbox("Expert Discussion", value=True, key="export_include_discussion")
        include_citations = st.checkbox("Citation List", value=True, key="export_include_citations")
        include_metadata = st.checkbox("Metadata", value=False, key="export_include_metadata")

    st.markdown("---")

    # Generate and download
    if st.button("Generate Export", type="primary", use_container_width=True):
        with st.spinner("Generating export..."):
            content = _generate_export_content(
                export_format=export_format,
                export_scope=export_scope,
                include_question=include_question,
                include_evidence=include_evidence,
                include_recommendation=include_recommendation,
                include_discussion=include_discussion,
                include_citations=include_citations,
                include_metadata=include_metadata,
                research_result=research_result,
                expert_discussion=expert_discussion,
                search_results=search_results,
                cdp_sections=cdp_sections
            )

            if content:
                st.session_state.export_content = content
                st.success("Export generated!")

    # Display preview and download button
    if st.session_state.get('export_content'):
        content = st.session_state.export_content

        st.markdown("### Preview")
        with st.expander("View Export Content", expanded=True):
            if export_format == "JSON":
                st.json(content if isinstance(content, dict) else {"content": content})
            else:
                st.markdown(content[:5000] + "..." if len(content) > 5000 else content)

        # Download button
        if export_format == "Markdown":
            st.download_button(
                label="Download Markdown",
                data=content,
                file_name=f"gdg_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                mime="text/markdown",
                use_container_width=True
            )
        elif export_format == "Word Document":
            # Generate Word document
            docx_buffer = _generate_word_document(content)
            if docx_buffer:
                st.download_button(
                    label="Download Word Document",
                    data=docx_buffer,
                    file_name=f"gdg_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        elif export_format == "JSON":
            import json
            json_str = json.dumps(content if isinstance(content, dict) else {"content": content}, indent=2)
            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name=f"gdg_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )


def _render_content_summary(research_result, expert_discussion, search_results, cdp_sections):
    """Render a summary of available content for export."""
    st.markdown("### Available Content")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if research_result:
            st.metric("Research Result", "Available")
        else:
            st.metric("Research Result", "None")

    with col2:
        rounds = len(expert_discussion) if expert_discussion else 0
        st.metric("Discussion Rounds", rounds)

    with col3:
        citations = search_results.get('citations', []) if search_results else []
        st.metric("Papers Found", len(citations))

    with col4:
        sections = len(cdp_sections) if cdp_sections else 0
        st.metric("Guideline Sections", sections)


def _get_available_scopes(research_result, expert_discussion, search_results, cdp_sections) -> List[str]:
    """Get available export scopes based on available data."""
    scopes = []

    if research_result:
        scopes.append("Latest Research Result")
    if expert_discussion:
        scopes.append("Full Discussion History")
    if search_results and search_results.get('citations'):
        scopes.append("Evidence Library")
    if cdp_sections:
        scopes.append("Guideline Workspace")

    if not scopes:
        scopes.append("No content available")

    return scopes


def _generate_export_content(
    export_format: str,
    export_scope: str,
    include_question: bool,
    include_evidence: bool,
    include_recommendation: bool,
    include_discussion: bool,
    include_citations: bool,
    include_metadata: bool,
    research_result: Optional[Dict],
    expert_discussion: Dict,
    search_results: Dict,
    cdp_sections: Dict
) -> str:
    """Generate export content based on selections."""

    sections = []

    # Header
    sections.append(f"# Palliative Surgery GDG Export")
    sections.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    sections.append(f"**Scope:** {export_scope}")
    sections.append("")

    # Clinical Question
    if include_question:
        question = None
        if research_result:
            question = research_result.get('question')
        elif expert_discussion:
            # Try to get from session state
            question = st.session_state.get('current_question') or st.session_state.get('expert_clinical_question')

        if question:
            sections.append("## Clinical Question")
            sections.append(f"> {question}")
            sections.append("")

    # Recommendation
    if include_recommendation and research_result:
        sections.append("## Recommendation")
        recommendation = research_result.get('recommendation', 'No recommendation available')
        confidence = research_result.get('confidence', 'Unknown')
        sections.append(f"**Confidence:** {confidence}")
        sections.append("")
        sections.append(recommendation)
        sections.append("")

        # Key findings
        key_findings = research_result.get('key_findings', [])
        if key_findings:
            sections.append("### Key Findings")
            for finding in key_findings:
                sections.append(f"- {finding}")
            sections.append("")

    # Expert Discussion
    if include_discussion and expert_discussion:
        sections.append("## Expert Discussion")
        for round_num in sorted(expert_discussion.keys()):
            round_data = expert_discussion[round_num]
            sections.append(f"### Round {round_num}")
            for expert_name, response in round_data.items():
                content = response.get('content', '') if isinstance(response, dict) else str(response)
                sections.append(f"#### {expert_name}")
                sections.append(content[:2000] + "..." if len(content) > 2000 else content)
                sections.append("")

    # Evidence Table
    if include_evidence and search_results:
        citations = search_results.get('citations', [])
        if citations:
            sections.append("## Evidence Summary")
            sections.append(f"**Papers found:** {len(citations)}")
            sections.append("")
            sections.append("| PMID | Title | Year | Journal |")
            sections.append("|------|-------|------|---------|")
            for cit in citations[:20]:  # Limit to 20 for readability
                pmid = cit.pmid if hasattr(cit, 'pmid') else cit.get('pmid', 'N/A')
                title = cit.title if hasattr(cit, 'title') else cit.get('title', 'Untitled')
                year = cit.year if hasattr(cit, 'year') else cit.get('year', 'N/A')
                journal = cit.journal if hasattr(cit, 'journal') else cit.get('journal', 'N/A')
                # Truncate title for table
                title_short = title[:60] + "..." if len(title) > 60 else title
                sections.append(f"| {pmid} | {title_short} | {year} | {journal} |")
            sections.append("")

    # Full Citation List
    if include_citations and search_results:
        citations = search_results.get('citations', [])
        if citations:
            sections.append("## References")
            for i, cit in enumerate(citations, 1):
                pmid = cit.pmid if hasattr(cit, 'pmid') else cit.get('pmid', '')
                title = cit.title if hasattr(cit, 'title') else cit.get('title', 'Untitled')
                authors = cit.authors if hasattr(cit, 'authors') else cit.get('authors', [])
                year = cit.year if hasattr(cit, 'year') else cit.get('year', '')
                journal = cit.journal if hasattr(cit, 'journal') else cit.get('journal', '')

                # Format authors
                if isinstance(authors, list):
                    if len(authors) > 3:
                        author_str = f"{authors[0]} et al."
                    else:
                        author_str = ", ".join(authors)
                else:
                    author_str = str(authors)

                sections.append(f"{i}. {author_str}. {title}. *{journal}*. {year}. PMID: {pmid}")
            sections.append("")

    # Guideline Workspace Sections
    if export_scope == "Guideline Workspace" and cdp_sections:
        sections.append("## Guideline Sections")
        for section_name, section_content in cdp_sections.items():
            sections.append(f"### {section_name}")
            if isinstance(section_content, dict):
                for key, value in section_content.items():
                    sections.append(f"**{key}:** {value}")
            else:
                sections.append(str(section_content))
            sections.append("")

    # Metadata
    if include_metadata:
        sections.append("## Metadata")
        sections.append(f"- **Export Date:** {datetime.now().isoformat()}")
        sections.append(f"- **Project:** {st.session_state.get('current_project_name', 'Unknown')}")
        sections.append(f"- **Model:** {settings.EXPERT_MODEL}")
        if research_result and research_result.get('metadata'):
            for key, value in research_result['metadata'].items():
                sections.append(f"- **{key}:** {value}")
        sections.append("")

    # Footer
    sections.append("---")
    sections.append("*Generated by Palliative Surgery GDG v1.0*")

    return "\n".join(sections)


def _generate_word_document(markdown_content: str) -> Optional[BytesIO]:
    """Convert markdown content to Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Parse markdown and add to document
        lines = markdown_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Headers
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=2)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=3)
            # Blockquotes
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:])
                p.style = 'Quote'
            # List items
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            # Numbered lists
            elif line[0].isdigit() and '. ' in line[:4]:
                idx = line.index('. ')
                p = doc.add_paragraph(line[idx+2:], style='List Number')
            # Table rows (skip table formatting markers)
            elif line.startswith('|') and not line.startswith('|---'):
                # Simple table handling - add as paragraph
                cells = [c.strip() for c in line.split('|')[1:-1]]
                p = doc.add_paragraph(' | '.join(cells))
            # Horizontal rules
            elif line.startswith('---'):
                p = doc.add_paragraph('_' * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Regular paragraphs
            else:
                # Clean up markdown formatting
                clean_line = line.replace('**', '').replace('*', '').replace('`', '')
                p = doc.add_paragraph(clean_line)

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except ImportError:
        st.warning("python-docx not installed. Word export unavailable.")
        return None
    except Exception as e:
        st.error(f"Failed to generate Word document: {e}")
        return None
