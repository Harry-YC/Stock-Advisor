"""
Document Ingestion Module

Handles ingestion of documents from various sources:
- PDF files
- DOCX files
- Plain text (paste)
- URLs (web articles)

All documents are normalized to a common Document format for use
in expert panel discussions and chat.
"""

import re
import uuid
import hashlib
from datetime import datetime
from typing import Optional, Dict, Any, List, BinaryIO
from dataclasses import dataclass, field, asdict
from pathlib import Path

import requests

# Optional imports with graceful degradation
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


@dataclass
class Document:
    """
    Unified document representation for all source types.

    This is the common format used throughout the app for evidence
    from any source (PubMed, preprints, uploads, etc.)
    """
    id: str
    source_type: str  # "pubmed", "preprint", "pdf", "docx", "text", "url"
    title: str
    content: str  # Full text or abstract
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: str = ""  # ISO format

    # Optional fields for compatibility with Citation
    authors: List[str] = field(default_factory=list)
    year: str = ""
    journal: str = ""
    doi: str = ""
    pmid: str = ""

    def __post_init__(self):
        if not self.created_at:
            self.created_at = datetime.now().isoformat()
        if not self.id:
            self.id = self._generate_id()

    def _generate_id(self) -> str:
        """Generate a unique ID based on content hash"""
        content_hash = hashlib.md5(
            f"{self.title}{self.content[:500]}".encode()
        ).hexdigest()[:12]
        return f"doc_{content_hash}"

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization"""
        return asdict(self)

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "Document":
        """Create Document from dictionary"""
        return cls(**data)

    def get_context_string(self, max_length: int = 2000) -> str:
        """Format document for LLM context"""
        parts = [f"**{self.title}**"]

        if self.authors:
            authors_str = ", ".join(self.authors[:3])
            if len(self.authors) > 3:
                authors_str += " et al."
            parts.append(f"Authors: {authors_str}")

        if self.year:
            parts.append(f"Year: {self.year}")

        if self.journal:
            parts.append(f"Source: {self.journal}")
        elif self.source_type:
            parts.append(f"Source: {self.source_type}")

        if self.pmid:
            parts.append(f"PMID: {self.pmid}")

        parts.append("")

        content = self.content
        if len(content) > max_length:
            content = content[:max_length] + "..."
        parts.append(content)

        return "\n".join(parts)


class DocumentIngestion:
    """Handles ingestion of documents from various sources"""

    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Literature Review App)"
        })

    def ingest_pdf(
        self,
        file: BinaryIO,
        filename: str = "Uploaded PDF"
    ) -> Document:
        """
        Extract text and metadata from a PDF file using layout-aware parsing.
        
        Uses block-based extraction to correctly handle multi-column layouts
        typical in scientific papers.

        Args:
            file: File-like object containing PDF data
            filename: Original filename for title fallback

        Returns:
            Document object with extracted content
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError(
                "PyMuPDF is required for PDF ingestion. "
                "Install with: pip install PyMuPDF"
            )

        # Read PDF
        pdf_bytes = file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        # Extract metadata
        metadata = doc.metadata or {}
        title = metadata.get("title", "") or filename.replace(".pdf", "")
        author = metadata.get("author", "")

        # Parse authors
        authors = []
        if author:
            # Try common separators
            for sep in [";", ",", " and ", " & "]:
                if sep in author:
                    authors = [a.strip() for a in author.split(sep)]
                    break
            if not authors:
                authors = [author]

        # Extract text using blocks (layout-aware)
        text_parts = []
        for page in doc:
            # get_text("blocks") returns list of (x0, y0, x1, y1, text, block_no, block_type)
            blocks = page.get_text("blocks")
            
            # Sort blocks: primary by vertical position (y0), secondary by horizontal (x0)
            # This is a heuristic for standard 2-column papers. 
            # PyMuPDF's default "blocks" usually respects reading order, but explicit sorting can help.
            # Actually, standard logic is: usually top-down, left-right. 
            # But let's trust fitz default reading order for now, it's generally good for columns.
            # For strict column handling, we can sort.
            blocks.sort(key=lambda b: (b[1], b[0])) 
            
            for b in blocks:
                # b[4] is the text content
                # b[6] is block_type (0=text, 1=image)
                if b[6] == 0: 
                    text_parts.append(b[4].strip())

        full_text = "\n\n".join(text_parts)

        # Clean up text
        full_text = self._clean_text(full_text)

        doc.close()

        return Document(
            id="",  # Will be auto-generated
            source_type="pdf",
            title=title,
            content=full_text,
            authors=authors,
            metadata={
                "filename": filename,
                "page_count": len(text_parts),
                "pdf_metadata": metadata,
                "extraction_method": "layout_blocks"
            }
        )

    def ingest_docx(
        self,
        file: BinaryIO,
        filename: str = "Uploaded Document"
    ) -> Document:
        """
        Extract text from a DOCX file.

        Args:
            file: File-like object containing DOCX data
            filename: Original filename for title

        Returns:
            Document object with extracted content
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX ingestion. "
                "Install with: pip install python-docx"
            )

        doc = DocxDocument(file)

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)

        # Try to extract title from first heading or filename
        title = filename.replace(".docx", "").replace(".doc", "")
        for para in doc.paragraphs:
            if para.style and "heading" in para.style.name.lower():
                if para.text.strip():
                    title = para.text.strip()
                    break

        return Document(
            id="",
            source_type="docx",
            title=title,
            content=full_text,
            metadata={
                "filename": filename,
                "paragraph_count": len(paragraphs)
            }
        )

    def ingest_text(
        self,
        text: str,
        title: str = "Pasted Text",
        source_description: str = ""
    ) -> Document:
        """
        Create a document from pasted text.

        Args:
            text: The pasted text content
            title: Title for the document
            source_description: Optional description of the source

        Returns:
            Document object
        """
        # Clean up text
        cleaned_text = self._clean_text(text)

        # Try to extract title from first line if not provided
        if title == "Pasted Text" and cleaned_text:
            first_line = cleaned_text.split("\n")[0].strip()
            if len(first_line) < 200:
                title = first_line

        return Document(
            id="",
            source_type="text",
            title=title,
            content=cleaned_text,
            metadata={
                "source_description": source_description,
                "character_count": len(cleaned_text)
            }
        )

    def ingest_url(
        self,
        url: str,
        timeout: int = 30
    ) -> Document:
        """
        Fetch and extract content from a URL.

        Args:
            url: The URL to fetch
            timeout: Request timeout in seconds

        Returns:
            Document object with extracted content
        """
        if not BS4_AVAILABLE:
            raise ImportError(
                "BeautifulSoup is required for URL ingestion. "
                "Install with: pip install beautifulsoup4"
            )

        # Fetch URL
        response = self.session.get(url, timeout=timeout)
        response.raise_for_status()

        # Parse HTML
        soup = BeautifulSoup(response.text, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        # Extract title
        title = ""
        if soup.title:
            title = soup.title.string or ""
        if not title:
            h1 = soup.find("h1")
            if h1:
                title = h1.get_text(strip=True)
        if not title:
            title = url

        # Extract main content
        # Try common content containers
        content_element = None
        for selector in ["article", "main", ".content", "#content", ".post", ".article"]:
            content_element = soup.select_one(selector)
            if content_element:
                break

        if not content_element:
            content_element = soup.body or soup

        # Extract text
        text = content_element.get_text(separator="\n", strip=True)
        text = self._clean_text(text)

        # Try to extract metadata
        authors = []
        meta_author = soup.find("meta", {"name": "author"})
        if meta_author and meta_author.get("content"):
            authors = [meta_author["content"]]

        date = ""
        for meta_name in ["date", "pubdate", "publish_date", "article:published_time"]:
            meta_date = soup.find("meta", {"name": meta_name}) or soup.find("meta", {"property": meta_name})
            if meta_date and meta_date.get("content"):
                date = meta_date["content"][:10]  # YYYY-MM-DD
                break

        return Document(
            id="",
            source_type="url",
            title=title,
            content=text,
            authors=authors,
            year=date[:4] if date else "",
            metadata={
                "url": url,
                "fetched_at": datetime.now().isoformat()
            }
        )

    def ingest_txt(
        self,
        file: BinaryIO,
        filename: str = "Uploaded Text"
    ) -> Document:
        """
        Read a plain text file.

        Args:
            file: File-like object containing text
            filename: Original filename

        Returns:
            Document object
        """
        # Try different encodings
        content = None
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                file.seek(0)
                content = file.read().decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise ValueError("Could not decode text file")

        title = filename.replace(".txt", "")

        return Document(
            id="",
            source_type="txt",
            title=title,
            content=self._clean_text(content),
            metadata={"filename": filename}
        )

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        # Remove common artifacts
        text = re.sub(r'\x00', '', text)  # Null bytes
        text = re.sub(r'[\x0b\x0c]', '\n', text)  # Vertical tabs, form feeds

        # Strip leading/trailing whitespace
        text = text.strip()

        return text


def ingest_file(
    file: BinaryIO,
    filename: str
) -> Document:
    """
    Convenience function to ingest a file based on extension.

    Args:
        file: File-like object
        filename: Original filename (used to determine type)

    Returns:
        Document object
    """
    ingestion = DocumentIngestion()

    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return ingestion.ingest_pdf(file, filename)
    elif ext in [".docx", ".doc"]:
        return ingestion.ingest_docx(file, filename)
    elif ext == ".txt":
        return ingestion.ingest_txt(file, filename)
    else:
        # Try as plain text
        return ingestion.ingest_txt(file, filename)


def documents_to_context(
    documents: List[Document],
    max_total_length: int = 50000
) -> str:
    """
    Format multiple documents as context for LLM prompts.

    Args:
        documents: List of Document objects
        max_total_length: Maximum total character length

    Returns:
        Formatted context string
    """
    if not documents:
        return ""

    parts = ["=" * 60, "EVIDENCE CORPUS", "=" * 60, ""]

    # Calculate per-document limit
    per_doc_limit = max_total_length // len(documents)

    for i, doc in enumerate(documents, 1):
        parts.append(f"[Document {i}]")
        parts.append(doc.get_context_string(max_length=per_doc_limit))
        parts.append("")
        parts.append("-" * 40)
        parts.append("")

    return "\n".join(parts)
