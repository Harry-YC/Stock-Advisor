"""
Guideline Workspace Component

Manages Guideline sections for Palliative Surgery GDG:
- View and edit guideline recommendation sections
- Track section status (draft/reviewed/approved)
- Preview complete guideline document
- Export to Word/PDF
"""

import streamlit as st
from datetime import datetime
from typing import Dict, List, Optional
import io

from core.question_templates import CDP_SECTION_MAPPING, QUESTION_TYPES


# Define canonical section order for Guideline
GUIDELINE_SECTION_ORDER = [
    "executive_summary",
    "patient_selection",
    "comparative_effectiveness",
    "symptom_management",
    "prognosis_outcomes",
    "ethical_considerations",
    "implementation_resources",
    "custom_section"
]

# Legacy alias
CDP_SECTION_ORDER = GUIDELINE_SECTION_ORDER

STATUS_COLORS = {
    "draft": ("#ffc107", "#fff3cd"),
    "reviewed": ("#17a2b8", "#d1ecf1"),
    "approved": ("#28a745", "#d4edda")
}


def render_cdp_workspace():
    """
    Render the Guideline Workspace view.

    Shows all saved guideline sections with editing and export capabilities.
    """
    st.title("Guideline Workspace")

    # Get CDP state
    if 'cdp_sections' not in st.session_state:
        st.session_state.cdp_sections = {}
        
    # Auto-Load from DB if session is empty but project is selected
    if not st.session_state.cdp_sections and st.session_state.get('current_project_id'):
        try:
            cdp_dao = st.session_state.get('cdp_dao')
            if cdp_dao:
                saved_sections = cdp_dao.get_cdp(st.session_state.current_project_id)
                if saved_sections:
                    st.session_state.cdp_sections = saved_sections
                    st.session_state.cdp_last_modified = datetime.now().isoformat()
        except Exception as e:
            st.warning(f"Failed to load saved CDP: {e}")

    cdp_sections = st.session_state.cdp_sections
    cdp_project_name = st.session_state.get('cdp_project_name', '')

    # Header with project info
    col1, col2 = st.columns([3, 1])
    with col1:
        new_name = st.text_input(
            "Guideline Name",
            value=cdp_project_name,
            placeholder="Enter guideline topic (e.g., Malignant Bowel Obstruction)...",
            key="cdp_project_name_input"
        )
        if new_name != cdp_project_name:
            st.session_state.cdp_project_name = new_name
            # Auto-save works via sections, but we should probably save name too in future.
            # For now, name is session-only until associated with sections.

    with col2:
        st.markdown("")  # Spacing
        st.markdown("")
        if st.button("Clear All Sections", type="secondary"):
            st.session_state.cdp_sections = {}
            _save_cdp_to_db()  # Persist clear
            st.rerun()

    # Show last modified
    last_modified = st.session_state.get('cdp_last_modified')
    if last_modified:
        st.caption(f"Last modified: {last_modified[:19].replace('T', ' ')}")

    st.markdown("---")

    # Check if any sections exist
    if not cdp_sections:
        st.info("""
        **No guideline sections yet.**

        Use the conversational interface to research questions and click "Add to Guideline"
        to build your clinical guideline section by section.

        **Quick Start:**
        1. Ask a clinical question (e.g., "When is surgical bypass preferred over stenting for malignant bowel obstruction?")
        2. Review the synthesized recommendation
        3. Click "Add to Guideline" to save it as a section
        4. Return here to view, edit, and export your guideline
        """)
        return

    # CDP Progress indicator
    total_sections = len(cdp_sections)
    approved_count = sum(1 for s in cdp_sections.values() if s.get('status') == 'approved')
    reviewed_count = sum(1 for s in cdp_sections.values() if s.get('status') == 'reviewed')

    st.progress(approved_count / total_sections if total_sections > 0 else 0)
    st.caption(f"{approved_count} approved / {reviewed_count} reviewed / {total_sections} total sections")

    # Action buttons row
    col_actions = st.columns([1, 1, 1, 2])
    with col_actions[0]:
        if st.button("Preview Guideline", type="primary"):
            st.session_state.cdp_preview_mode = True
    with col_actions[1]:
        if st.button("Export to Word"):
            docx_bytes = export_cdp_to_docx(cdp_sections, cdp_project_name)
            if docx_bytes:
                st.download_button(
                    label="Download .docx",
                    data=docx_bytes,
                    file_name=f"Guideline_{cdp_project_name or 'draft'}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    with col_actions[2]:
        if st.button("Export to Markdown"):
            md_content = export_cdp_to_markdown(cdp_sections, cdp_project_name)
            st.download_button(
                label="Download .md",
                data=md_content,
                file_name=f"Guideline_{cdp_project_name or 'draft'}.md",
                mime="text/markdown"
            )

    st.markdown("---")

    # Preview mode
    if st.session_state.get('cdp_preview_mode', False):
        render_cdp_preview(cdp_sections, cdp_project_name)
        if st.button("Exit Preview"):
            st.session_state.cdp_preview_mode = False
            st.rerun()
        return

    # Section list
    st.subheader("Guideline Sections")

    # Sort sections by canonical order
    sorted_sections = sorted(
        cdp_sections.items(),
        key=lambda x: CDP_SECTION_ORDER.index(x[0]) if x[0] in CDP_SECTION_ORDER else 999
    )

    for section_key, section in sorted_sections:
        render_section_card(section_key, section)


def render_section_card(section_key: str, section: Dict):
    """Render a single guideline section card with edit/delete controls."""
    title = section.get('title', section_key.replace('_', ' ').title())
    status = section.get('status', 'draft')
    confidence = section.get('confidence', 'MEDIUM')
    question = section.get('source_question', '')
    created_at = section.get('created_at', '')

    # Status colors
    status_color, status_bg = STATUS_COLORS.get(status, ("#666", "#f5f5f5"))

    # Section header
    with st.expander(f"{title}", expanded=False):
        # Status and metadata row
        col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
        with col1:
            st.markdown(f"""
            <span style="background: {status_bg}; color: {status_color};
                        padding: 0.25rem 0.5rem; border-radius: 0.25rem;
                        font-size: 0.8rem; font-weight: bold;">
                {status.upper()}
            </span>
            """, unsafe_allow_html=True)
        with col2:
            st.caption(f"Confidence: {confidence}")
        with col3:
            if created_at:
                st.caption(f"Created: {created_at[:10]}")
        with col4:
            # Status change dropdown
            new_status = st.selectbox(
                "Change status",
                ["draft", "reviewed", "approved"],
                index=["draft", "reviewed", "approved"].index(status),
                key=f"status_{section_key}",
                label_visibility="collapsed"
            )
            if new_status != status:
                st.session_state.cdp_sections[section_key]['status'] = new_status
                st.session_state.cdp_sections[section_key]['last_edited'] = datetime.now().isoformat()
                _save_cdp_to_db()  # Persist status change
                st.rerun()

        # Source question
        if question:
            st.caption(f"**Source question:** {question[:100]}...")

        # Content preview/edit
        content = section.get('content', '')

        # Edit mode toggle
        edit_key = f"edit_mode_{section_key}"
        if edit_key not in st.session_state:
            st.session_state[edit_key] = False

        col_edit1, col_edit2 = st.columns([1, 5])
        with col_edit1:
            if st.button("Edit" if not st.session_state[edit_key] else "Save", key=f"edit_btn_{section_key}"):
                if st.session_state[edit_key]:
                    # Save mode - get edited content
                    edited_content = st.session_state.get(f"content_edit_{section_key}", content)
                    st.session_state.cdp_sections[section_key]['content'] = edited_content
                    st.session_state.cdp_sections[section_key]['last_edited'] = datetime.now().isoformat()
                    _save_cdp_to_db()  # Persist content edit
                st.session_state[edit_key] = not st.session_state[edit_key]
                st.rerun()

        if st.session_state[edit_key]:
            # Edit mode
            edited = st.text_area(
                "Edit section content",
                value=content,
                height=400,
                key=f"content_edit_{section_key}",
                label_visibility="collapsed"
            )
        else:
            # View mode
            st.markdown(content)

        # Key findings
        findings = section.get('key_findings', [])
        if findings:
            st.markdown("**Key Findings:**")
            for finding in findings:
                st.markdown(f"- {finding}")

        # Expert contributors
        experts = section.get('expert_contributors', [])
        if experts:
            st.caption(f"**Expert contributors:** {', '.join(experts)}")

        # Citations
        citations = section.get('citations', [])
        if citations:
            with st.expander(f"Citations ({len(citations)})"):
                for cit in citations:
                    pmid = cit.get('pmid', '')
                    title = cit.get('title', 'Untitled')
                    st.caption(f"- {title[:60]}... (PMID: {pmid})")

        # Delete button
        st.markdown("---")
        if st.button("Delete Section", key=f"delete_{section_key}", type="secondary"):
            del st.session_state.cdp_sections[section_key]
            _save_cdp_to_db()  # Persist deletion
            st.rerun()


def _save_cdp_to_db():
    """Helper to persist guideline state to database."""
    if st.session_state.get('current_project_id') and st.session_state.get('cdp_dao'):
        try:
            st.session_state.cdp_dao.save_cdp(
                st.session_state.current_project_id,
                st.session_state.cdp_sections
            )
        except Exception as e:
            st.warning(f"Auto-save failed: {e}")


def render_cdp_preview(cdp_sections: Dict, project_name: str):
    """Render a full preview of the guideline document."""
    st.markdown(f"""
    # Palliative Surgery Clinical Guideline
    ## {project_name or '[Guideline Topic]'}

    *Generated: {datetime.now().strftime('%Y-%m-%d')}*

    ---
    """)

    # Sort and render sections
    sorted_sections = sorted(
        cdp_sections.items(),
        key=lambda x: CDP_SECTION_ORDER.index(x[0]) if x[0] in CDP_SECTION_ORDER else 999
    )

    for section_key, section in sorted_sections:
        title = section.get('title', section_key.replace('_', ' ').title())
        content = section.get('content', '')
        status = section.get('status', 'draft')

        # Section header with status badge
        status_emoji = {"draft": "", "reviewed": "", "approved": ""}
        st.markdown(f"## {title} {status_emoji.get(status, '')}")

        # Content
        st.markdown(content)

        st.markdown("---")


def export_cdp_to_markdown(cdp_sections: Dict, project_name: str) -> str:
    """Export guideline to Markdown format."""
    lines = [
        f"# Palliative Surgery Clinical Guideline",
        f"## {project_name or '[Guideline Topic]'}",
        "",
        f"*Generated: {datetime.now().strftime('%Y-%m-%d')} by Palliative Surgery GDG*",
        "",
        "---",
        ""
    ]

    # Sort sections
    sorted_sections = sorted(
        cdp_sections.items(),
        key=lambda x: CDP_SECTION_ORDER.index(x[0]) if x[0] in CDP_SECTION_ORDER else 999
    )

    for section_key, section in sorted_sections:
        title = section.get('title', section_key.replace('_', ' ').title())
        content = section.get('content', '')

        lines.append(f"## {title}")
        lines.append("")
        lines.append(content)
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def export_cdp_to_docx(cdp_sections: Dict, project_name: str) -> Optional[bytes]:
    """
    Export guideline to Word document format.

    Returns:
        Bytes of the docx file, or None if python-docx is not available
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        st.warning("python-docx not installed. Run: pip install python-docx")
        return None

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading('Palliative Surgery Clinical Guideline', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph(project_name or '[Guideline Topic]')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')} by Palliative Surgery GDG")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Sort sections
    sorted_sections = sorted(
        cdp_sections.items(),
        key=lambda x: CDP_SECTION_ORDER.index(x[0]) if x[0] in CDP_SECTION_ORDER else 999
    )

    for section_key, section in sorted_sections:
        title = section.get('title', section_key.replace('_', ' ').title())
        content = section.get('content', '')

        # Section heading
        doc.add_heading(title, level=1)

        # Content - split by paragraphs
        for para in content.split('\n\n'):
            if para.strip():
                # Handle markdown headers within content
                if para.startswith('## '):
                    doc.add_heading(para[3:], level=2)
                elif para.startswith('### '):
                    doc.add_heading(para[4:], level=3)
                else:
                    doc.add_paragraph(para.strip())

        doc.add_paragraph()  # Spacing between sections

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def get_cdp_section_count() -> int:
    """Get count of guideline sections for sidebar display."""
    return len(st.session_state.get('cdp_sections', {}))


def has_cdp_sections() -> bool:
    """Check if any guideline sections exist."""
    return len(st.session_state.get('cdp_sections', {}))
